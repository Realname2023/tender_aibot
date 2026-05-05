import zipfile
import subprocess
import tempfile
import os
import base64
import re
from io import BytesIO
from dataclasses import dataclass, field
from typing import Optional

import pymupdf                          # fitz — PDF
from docx import Document as DocxDoc   # python-docx — DOCX
import mammoth                          # DOC/DOCX → текст (без Spire)
import openpyxl                         # XLSX
import xlrd                             # XLS
import pycld2 as cld2                   # определение языка


# Уникальные казахские буквы которых НЕТ в русском
_KK_ONLY = set('әіңғүұқөһӘІҢҒҮҰҚӨҺ')
# Уникальные русские буквы которых НЕТ в казахском
_RU_ONLY = set('ёъыЁЪЫ')


def detect_lang(text: str) -> str:
    """
    Определяет язык текста.
    Возвращает: 'ru', 'kk', 'other', 'unknown'

    Стратегия:
    1. Уникальные буквы — самый надёжный способ для kk/ru
    2. pycld2 — если уникальных букв нет
    3. Доля кириллицы — запасной вариант
    """
    if not text or len(text.strip()) < 5:
        return 'unknown'

    # Шаг 1: уникальные буквы
    kk_chars = sum(1 for c in text if c in _KK_ONLY)
    ru_chars = sum(1 for c in text if c in _RU_ONLY)

    if kk_chars > 0 and ru_chars == 0:
        return 'kk'
    if ru_chars > 0 and kk_chars == 0:
        return 'ru'
    if kk_chars > ru_chars:
        return 'kk'
    if ru_chars > kk_chars:
        return 'ru'

    # Шаг 2: pycld2
    try:
        is_reliable, _, details = cld2.detect(text)
        if is_reliable and details:
            lang = details[0][1]
            if lang == 'ru':
                return 'ru'
            if lang in ('kk', 'uz', 'az'):  # cld2 путает казахский с узбекским/азербайджанским
                return 'kk'
            if lang != 'un':
                return 'other'
    except Exception:
        pass

    # Шаг 3: есть ли вообще кириллица
    cyrillic = len(re.findall(r'[а-яёА-ЯЁәіңғүұқөһӘІҢҒҮҰҚӨҺ]', text))
    if cyrillic / max(len(text), 1) > 0.3:
        return 'unknown'  # кириллица есть но язык не определён — оставляем

    return 'other'


def filter_by_lang(text: str, target_lang: str = 'ru') -> str:
    """
    Фильтрует текст по языку — оставляет только абзацы на нужном языке.
    Абзацы с 'unknown' (цифры, даты, термины) тоже оставляем.
    target_lang: 'ru' или 'kk'
    """
    if not text:
        return ''

    filtered = []
    for para in text.split('\n'):
        stripped = para.strip()
        if not stripped:
            filtered.append('')
            continue
        lang = detect_lang(stripped)
        if lang in (target_lang, 'unknown'):
            filtered.append(para)

    return '\n'.join(filtered).strip()


# ---------------------------------------------------------------------------
# Структура результата
# ---------------------------------------------------------------------------

@dataclass
class DocResult:
    """
    Результат извлечения документа.
    text       — текст документа (отфильтрованный по языку)
    tables     — список таблиц; каждая таблица = список строк; строка = список ячеек
    images_b64 — список картинок в base64 (JPEG/PNG)
    file_name  — имя файла
    error      — сообщение об ошибке если что-то пошло не так
    """
    file_name: str = ""
    text: str = ""
    tables: list = field(default_factory=list)
    images_b64: list = field(default_factory=list)
    error: Optional[str] = None

    def __repr__(self):
        return (
            f"<DocResult: {self.file_name} | "
            f"текст={len(self.text)} симв. | "
            f"таблиц={len(self.tables)} | "
            f"картинок={len(self.images_b64)}>"
        )

    def to_agent_text(self) -> str:
        """Текстовое представление для LLM агента."""
        parts = []

        if self.error:
            return f"[ОШИБКА] {self.error}"

        if self.text.strip():
            parts.append(f"=== ТЕКСТ ===\n{self.text.strip()}")

        for i, table in enumerate(self.tables):
            lines = [f"\n=== ТАБЛИЦА {i + 1} ==="]
            for row in table:
                lines.append(" | ".join(str(c) for c in row))
            parts.append("\n".join(lines))

        if self.images_b64:
            parts.append(f"\n[В документе {len(self.images_b64)} изображений.]")

        return "\n\n".join(parts) if parts else "[Документ пустой или не содержит читаемого текста]"

    def to_agent_images(self) -> list[dict]:
        """Картинки в формате для vision LLM."""
        return [
            {
                "type": "image_url",
                "image_url": {"url": f"data:image/jpeg;base64,{b64}"}
            }
            for b64 in self.images_b64
        ]


# ---------------------------------------------------------------------------
# Утилиты
# ---------------------------------------------------------------------------

def _pix_to_b64(pix: pymupdf.Pixmap) -> str:
    """Pixmap → base64 JPEG строка."""
    if pix.alpha:
        pix = pymupdf.Pixmap(pymupdf.csRGB, pix)
    buf = BytesIO(pix.tobytes("jpeg"))
    return base64.b64encode(buf.getvalue()).decode()


def _img_bytes_to_b64(data: bytes) -> str:
    return base64.b64encode(data).decode()


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def extract_from_pdf(file_name: str, pdf_bytes: bytes, lang: str = 'ru') -> DocResult:
    """
    Извлекает текст и картинки из PDF.
    - Текстовые страницы → берём текст, фильтруем по языку
    - Страницы без текста (сканы) → скриншот страницы (72 dpi)
    - Встроенные изображения → собираем
    lang: 'ru' или 'kk' — язык для фильтрации текста
    """
    result = DocResult(file_name=file_name)
    texts = []
    images_b64 = []

    try:
        with pymupdf.open(stream=pdf_bytes, filetype="pdf") as doc:
            for page_num, page in enumerate(doc):
                text = page.get_text().strip()
                raw_images = page.get_images(full=True)

                if text:
                    filtered = filter_by_lang(text, target_lang=lang)
                    if filtered:
                        texts.append(f"[Стр. {page_num + 1}]\n{filtered}")

                # Встроенные изображения страницы
                for img_info in raw_images:
                    xref = img_info[0]
                    try:
                        base_img = doc.extract_image(xref)
                        images_b64.append(_img_bytes_to_b64(base_img["image"]))
                    except Exception:
                        pass

                # Страница без текста — скан
                if not text:
                    pix = page.get_pixmap(dpi=72)  # 72 dpi — легче в 4 раза чем 150
                    images_b64.append(_pix_to_b64(pix))

        result.text = "\n\n".join(texts)
        result.images_b64 = images_b64

    except Exception as e:
        result.error = f"PDF: {e}"

    return result


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def extract_from_docx(file_name: str, docx_bytes: bytes, lang: str = 'ru') -> DocResult:
    """
    Извлекает текст, таблицы и картинки из DOCX.
    lang: 'ru' или 'kk' — язык для фильтрации текста
    """
    result = DocResult(file_name=file_name)
    texts = []
    tables = []
    images_b64 = []

    try:
        doc = DocxDoc(BytesIO(docx_bytes))

        # Текст параграфов — фильтруем по языку
        for para in doc.paragraphs:
            t = para.text.strip()
            if t and detect_lang(t) in (lang, 'unknown'):
                texts.append(t)

        # Таблицы — берём все (часто двуязычные)
        for table in doc.tables:
            tbl = []
            for row in table.rows:
                tbl.append([cell.text.strip() for cell in row.cells])
            tables.append(tbl)

        # Изображения из ZIP-структуры DOCX
        with zipfile.ZipFile(BytesIO(docx_bytes)) as z:
            for name in z.namelist():
                if name.startswith("word/media/") and name.lower().endswith(
                        (".png", ".jpg", ".jpeg", ".gif", ".bmp")):
                    images_b64.append(_img_bytes_to_b64(z.read(name)))

        result.text = "\n".join(texts)
        result.tables = tables
        result.images_b64 = images_b64

    except Exception as e:
        result.error = f"DOCX: {e}"

    return result


# ---------------------------------------------------------------------------
# DOC (старый бинарный формат) — через mammoth
# ---------------------------------------------------------------------------

def extract_from_doc(file_name: str, doc_bytes: bytes, lang: str = 'ru') -> DocResult:
    """
    Извлекает текст из DOC через mammoth.
    Если mammoth не справился — пробуем LibreOffice.
    lang: 'ru' или 'kk' — язык для фильтрации текста
    """
    result = DocResult(file_name=file_name)

    # --- Попытка 1: mammoth ---
    try:
        output = mammoth.extract_raw_text(BytesIO(doc_bytes))
        text = output.value.strip()
        if text:
            result.text = filter_by_lang(text, target_lang=lang)
            return result
    except Exception:
        pass

    # --- Попытка 2: LibreOffice ---
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = os.path.join(tmpdir, file_name)
            with open(doc_path, "wb") as f:
                f.write(doc_bytes)

            subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "docx",
                 "--outdir", tmpdir, doc_path],
                capture_output=True, timeout=30
            )

            docx_path = doc_path.replace(".doc", ".docx")
            if os.path.exists(docx_path):
                with open(docx_path, "rb") as f:
                    return extract_from_docx(file_name, f.read(), lang=lang)

    except (subprocess.TimeoutExpired, FileNotFoundError):
        pass
    except Exception as e:
        result.error = f"DOC LibreOffice: {e}"

    if not result.text and not result.error:
        result.error = "DOC: не удалось извлечь текст (нет mammoth/libreoffice)"

    return result


# ---------------------------------------------------------------------------
# XLSX
# ---------------------------------------------------------------------------

def extract_from_xlsx(file_name: str, xlsx_bytes: bytes) -> DocResult:
    """Извлекает таблицы из XLSX через openpyxl."""
    result = DocResult(file_name=file_name)
    tables = []

    try:
        wb = openpyxl.load_workbook(BytesIO(xlsx_bytes), data_only=True)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            table = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(c.strip() for c in cells):
                    table.append(cells)
            if table:
                tables.append(table)

        result.tables = tables

    except Exception as e:
        result.error = f"XLSX: {e}"

    return result


# ---------------------------------------------------------------------------
# XLS
# ---------------------------------------------------------------------------

def extract_from_xls(file_name: str, xls_bytes: bytes) -> DocResult:
    """Извлекает таблицы из XLS через xlrd."""
    result = DocResult(file_name=file_name)
    tables = []

    try:
        wb = xlrd.open_workbook(file_contents=xls_bytes)
        for sheet in wb.sheets():
            table = []
            for row_idx in range(sheet.nrows):
                cells = [str(v) for v in sheet.row_values(row_idx)]
                if any(c.strip() for c in cells):
                    table.append(cells)
            if table:
                tables.append(table)

        result.tables = tables

    except Exception as e:
        result.error = f"XLS: {e}"

    return result


# ---------------------------------------------------------------------------
# Архивы ZIP / RAR
# ---------------------------------------------------------------------------

def extract_from_zip(file_name: str, zip_bytes: bytes, lang: str = 'ru') -> list[DocResult]:
    """Извлекает все документы из ZIP архива."""
    results = []
    supported = (".pdf", ".doc", ".docx", ".xlsx", ".xls")

    try:
        with zipfile.ZipFile(BytesIO(zip_bytes)) as z:
            for info in z.infolist():
                name = info.filename
                if name.lower().endswith(supported) and not name.startswith("__"):
                    content = z.read(name)
                    results.append(_dispatch(name, content, lang=lang))
    except Exception as e:
        results.append(DocResult(file_name=file_name, error=f"ZIP: {e}"))

    return results


def extract_from_rar(file_name: str, rar_bytes: bytes, lang: str = 'ru') -> list[DocResult]:
    """
    Извлекает документы из RAR.
    Требует: sudo apt install unrar
    """
    results = []
    supported = (".pdf", ".doc", ".docx", ".xlsx", ".xls")

    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            rar_path = os.path.join(tmpdir, "archive.rar")
            with open(rar_path, "wb") as f:
                f.write(rar_bytes)

            proc = subprocess.run(
                ["unrar", "x", "-y", rar_path, tmpdir],
                capture_output=True, timeout=60
            )

            if proc.returncode != 0:
                err = proc.stderr.decode(errors="replace")
                return [DocResult(file_name=file_name, error=f"RAR unrar: {err}")]

            for root, _, files in os.walk(tmpdir):
                for fname in files:
                    if fname.lower().endswith(supported):
                        fpath = os.path.join(root, fname)
                        with open(fpath, "rb") as f:
                            content = f.read()
                        results.append(_dispatch(fname, content, lang=lang))

    except FileNotFoundError:
        results.append(DocResult(
            file_name=file_name,
            error="RAR: unrar не установлен. Выполни: sudo apt install unrar"
        ))
    except Exception as e:
        results.append(DocResult(file_name=file_name, error=f"RAR: {e}"))

    return results


# ---------------------------------------------------------------------------
# Диспетчер по расширению
# ---------------------------------------------------------------------------

def _dispatch(file_name: str, content: bytes, lang: str = 'ru') -> DocResult:
    ext = file_name.lower().rsplit(".", 1)[-1]
    if ext == "pdf":
        return extract_from_pdf(file_name, content, lang=lang)
    elif ext == "docx":
        return extract_from_docx(file_name, content, lang=lang)
    elif ext == "doc":
        return extract_from_doc(file_name, content, lang=lang)
    elif ext == "xlsx":
        return extract_from_xlsx(file_name, content)
    elif ext == "xls":
        return extract_from_xls(file_name, content)
    return DocResult(file_name=file_name, error=f"Формат .{ext} не поддерживается")


def extract_document(file_name: str, file_bytes: bytes, lang: str = 'ru') -> list[DocResult]:
    """
    Главная функция. Принимает имя файла и его байты.
    Возвращает список DocResult.
    lang: 'ru' или 'kk' — язык для фильтрации текста (по умолчанию русский)
    """
    ext = file_name.lower().rsplit(".", 1)[-1]
    if ext == "zip":
        return extract_from_zip(file_name, file_bytes, lang=lang)
    elif ext == "rar":
        return extract_from_rar(file_name, file_bytes, lang=lang)
    else:
        return [_dispatch(file_name, file_bytes, lang=lang)]
